import subprocess

import pytest

from cv_creator_reborn.opencode_client import (
    FreeUsageLimitError,
    ask_opencode,
)


class TestAskOpencode:
    def test_returns_response_on_success(self, mocker):
        completed = subprocess.CompletedProcess(
            args=["opencode", "run", "test"],
            returncode=0,
            stdout="Hello\n",
            stderr="",
        )
        mocker.patch("subprocess.run", return_value=completed)

        result = ask_opencode("Say hello")

        assert result == "Hello"

    def test_strips_ansi_and_metadata(self, mocker):
        ansi_output = (
            "\x1b[0m\n"
            "> build \xB7 big-pickle\n"
            "\x1b[0m\n"
            "The response text\n"
        )
        completed = subprocess.CompletedProcess(
            args=["opencode", "run", "test"],
            returncode=0,
            stdout=ansi_output,
            stderr="",
        )
        mocker.patch("subprocess.run", return_value=completed)

        result = ask_opencode("Test query")

        assert result == "The response text"

    def test_raises_free_limit_error_on_rate_limit(self, mocker):
        error_output = (
            'Error: 429: {"type":"error","error":'
            '{"type":"FreeUsageLimitError",'
            '"message":"Rate limit exceeded. Please try again later."}}\n'
        )
        completed = subprocess.CompletedProcess(
            args=["opencode", "run", "test"],
            returncode=1,
            stdout="",
            stderr=error_output,
        )
        mocker.patch("subprocess.run", return_value=completed)

        with pytest.raises(FreeUsageLimitError, match="Rate limit exceeded"):
            ask_opencode("Test query")

    def test_raises_free_limit_error_on_stderr_message(self, mocker):
        completed = subprocess.CompletedProcess(
            args=["opencode", "run", "test"],
            returncode=1,
            stdout="",
            stderr="FreeUsageLimitError: monthly limit reached",
        )
        mocker.patch("subprocess.run", return_value=completed)

        with pytest.raises(FreeUsageLimitError, match="monthly limit reached"):
            ask_opencode("Test query")

    def test_raises_runtime_error_on_unknown_failure(self, mocker):
        completed = subprocess.CompletedProcess(
            args=["opencode", "run", "test"],
            returncode=1,
            stdout="",
            stderr="Some other error",
        )
        mocker.patch("subprocess.run", return_value=completed)

        with pytest.raises(RuntimeError, match="opencode failed"):
            ask_opencode("Test query")


class TestDocxHandlerOpencode:
    def test_execute_opencode_prompts_replaces_tags(self, mocker, tmp_path):
        from docx import Document

        from cv_creator_reborn.documents import DocxHandler

        docx_path = tmp_path / "test.docx"
        doc = Document()
        doc.add_paragraph("Hello <CHAT_GPT>write a greeting</CHAT_GPT> world")
        doc.save(str(docx_path))

        handler = DocxHandler()
        handler.open_document(str(docx_path))

        mock_response = "Good day"
        completed = subprocess.CompletedProcess(
            args=["opencode", "run", "test"],
            returncode=0,
            stdout=mock_response,
            stderr="",
        )
        mocker.patch("subprocess.run", return_value=completed)

        handler.execute_opencode_prompts("Some job description")

        handler.save_document(str(tmp_path / "out.docx"))
        handler.close_document()

        result_doc = Document(str(tmp_path / "out.docx"))
        text = " ".join(p.text for p in result_doc.paragraphs)
        assert "Good day" in text
        assert "<CHAT_GPT>" not in text

    def test_free_limit_during_prompt_execution_raises_error(
        self, mocker, tmp_path
    ):
        from docx import Document

        from cv_creator_reborn.documents import DocxHandler

        docx_path = tmp_path / "test.docx"
        doc = Document()
        doc.add_paragraph("<CHAT_GPT>write a greeting</CHAT_GPT>")
        doc.save(str(docx_path))

        handler = DocxHandler()
        handler.open_document(str(docx_path))

        error_output = (
            'Error: 429: {"type":"error","error":'
            '{"type":"FreeUsageLimitError",'
            '"message":"Rate limit exceeded."}}\n'
        )
        completed = subprocess.CompletedProcess(
            args=["opencode", "run", "test"],
            returncode=1,
            stdout="",
            stderr=error_output,
        )
        mocker.patch("subprocess.run", return_value=completed)

        with pytest.raises(FreeUsageLimitError):
            handler.execute_opencode_prompts("Some job description")

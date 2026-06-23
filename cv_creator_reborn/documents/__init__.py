import re
from abc import ABC, abstractmethod

from cv_creator_reborn.opencode_client import ask_opencode


class DocumentHandler(ABC):
    @abstractmethod
    def open_document(self, file_path):
        pass

    @abstractmethod
    def update_document(self, key, value):
        pass

    @abstractmethod
    def save_document(self, output_path=None):
        pass

    @abstractmethod
    def close_document(self):
        pass


class DocxHandler(DocumentHandler):
    def __init__(self):
        self.document = None

    def open_document(self, file_path):
        try:
            import docx

            self.document = docx.Document(file_path)
            return True
        except ImportError:
            raise ImportError(
                "python-docx library is required for .docx files. Install with: pip install python-docx"
            )
        except Exception as e:
            raise Exception(f"Error opening .docx file: {e}")

    def return_opencode_response(self, query, job_description):
        prompt = f"{query}: {job_description}"
        return ask_opencode(prompt)

    def execute_opencode_prompts(self, description):
        start_tag = "<AI>"
        end_tag = "</AI>"
        for paragraph in self.document.paragraphs:
            matches = re.findall(rf"{start_tag}(.*?){end_tag}", paragraph.text)
            if matches:
                for match in matches:
                    for tag in [start_tag, end_tag]:
                        match = match.replace(tag, "")

                    query_response = self.return_opencode_response(match, description)
                    placeholder_text = f"{start_tag}{match}{end_tag}"
                    paragraph.text = paragraph.text.replace(
                        placeholder_text, query_response, 1
                    )

    def update_document(self, key, value):
        if not self.document:
            raise ValueError("Document not opened")

        for paragraph in self.document.paragraphs:
            if key in paragraph.text:
                try:
                    runs = paragraph.runs
                    for i in range(len(runs)):
                        if key in runs[i].text:
                            text = runs[i].text.replace(key, value)
                            runs[i].text = text
                except (TypeError, AttributeError):
                    pass

        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if key in paragraph.text:
                            try:
                                runs = paragraph.runs
                                for i in range(len(runs)):
                                    if key in runs[i].text:
                                        text = runs[i].text.replace(key, value)
                                        runs[i].text = text
                            except (TypeError, AttributeError):
                                pass

    def save_document(self, output_path=None):
        if not self.document:
            raise ValueError("Document not opened")

        if output_path:
            self.document.save(output_path)
        else:
            raise ValueError("Output path required for .docx files")

    def close_document(self):
        self.document = None

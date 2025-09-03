import os
from abc import ABC, abstractmethod
from pathlib import Path


class DocumentHandler(ABC):
    @abstractmethod
    def open_document(self, file_path):
        pass

    @abstractmethod
    def update_document(self, key, value):
        pass

    @abstractmethod
    def save_document(self, output_path=None):
        pass

    @abstractmethod
    def close_document(self):
        pass


class DocxHandler(DocumentHandler):
    def __init__(self):
        self.document = None

    def open_document(self, file_path):
        try:
            import docx

            self.document = docx.Document(file_path)
            return True
        except ImportError:
            raise ImportError(
                "python-docx library is required for .docx files. Install with: pip install python-docx"
            )
        except Exception as e:
            raise Exception(f"Error opening .docx file: {e}")

    def update_document(self, key, value):
        if not self.document:
            raise ValueError("Document not opened")

        for paragraph in self.document.paragraphs:
            if key in paragraph.text:
                try:
                    runs = paragraph.runs
                    for i in range(len(runs)):
                        if key in runs[i].text:
                            text = runs[i].text.replace(key, value)
                            runs[i].text = text
                except (TypeError, AttributeError):
                    pass

        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if key in paragraph.text:
                            try:
                                runs = paragraph.runs
                                for i in range(len(runs)):
                                    if key in runs[i].text:
                                        text = runs[i].text.replace(key, value)
                                        runs[i].text = text
                            except (TypeError, AttributeError):
                                pass

    def save_document(self, output_path=None):
        if not self.document:
            raise ValueError("Document not opened")

        if output_path:
            self.document.save(output_path)
        else:
            raise ValueError("Output path required for .docx files")

    def close_document(self):
        self.document = None


class OdtHandler(DocumentHandler):
    def __init__(self):
        raise NotImplementedError("Currently not replacing text correctly.")
        self.document = None
        self.file_path = None

    def open_document(self, file_path):
        try:
            from odf.opendocument import load

            self.document = load(file_path)
            self.file_path = file_path
            return True
        except ImportError:
            raise ImportError(
                "odfpy library is required for .odt files. Install with: pip install odfpy"
            )
        except Exception as e:
            raise Exception(f"Error opening .odt file: {e}")

    def update_document(self, key, value):
        if not self.document:
            raise ValueError("Document not opened")

        from odf import teletype, text

        paragraphs = self.document.getElementsByType(text.P)

        for paragraph in paragraphs:

            original_text = teletype.extractText(paragraph)

            if key in original_text:
                modified_text = original_text.replace(key, value)

                new_paragraph = text.P()

                original_style = paragraph.getAttribute("stylename")
                new_paragraph.setAttribute("stylename", original_style)

                new_paragraph.addText(modified_text)

                parent = paragraph.parentNode
                parent.insertBefore(new_paragraph, paragraph)
                parent.removeChild(paragraph)

    def save_document(self, output_path=None):
        if not self.document:
            raise ValueError("Document not opened")

        save_path = output_path if output_path else self.file_path
        self.document.save(save_path)

    def close_document(self):
        self.document = None
        self.file_path = None


class AgnosticDocument:
    def __init__(self):
        self.handler = None
        self.file_type = None

    def open_document(self, file_path):
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        extension = file_path.suffix.lower()

        if extension == ".docx":
            self.handler = DocxHandler()
            self.file_type = "docx"
        elif extension == ".odt":
            self.handler = OdtHandler()
            self.file_type = "odt"
        else:
            raise ValueError(
                f"Unsupported file format: {extension}. Supported formats: .docx, .odt"
            )

        return self.handler.open_document(file_path)

    def update_document(self, key, value):
        if not self.handler:
            raise ValueError("No document opened. Call open_document() first.")

        self.handler.update_document(key, value)

    def update_multiple(self, replacements):
        for key, value in replacements.items():
            self.update_document(key, value)

    def save_document(self, output_path):
        if not self.handler:
            raise ValueError("No document opened. Call open_document() first.")

        if os.path.exists(output_path):
            os.remove(output_path)

        self.handler.save_document(output_path)

    def close_document(self):
        if self.handler:
            self.handler.close_document()
        self.handler = None
        self.file_type = None

    def get_supported_formats(self):
        return [".docx", ".odt"]

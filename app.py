import requests
import os
import sys
import re
from docx import Document
import shutil
import time
import json
import random
import pyperclip
from selenium import webdriver
from selenium.webdriver.chrome.service import Service as ChromeService
from selenium.webdriver.common.by import By
from selenium.common.exceptions import TimeoutException
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from openai import OpenAI
import tkinter as tk
from tkinter import ttk, filedialog, NORMAL, DISABLED
import threading


class Seeker(object):
    def __init__(self, name, temp_path_dict, where, keywords, view_previous_jobs):
        self.name = name
        self.driver = None
        self.temp_path_dict = temp_path_dict
        self.where = where
        self.keywords = keywords
        self.view_previous_jobs = view_previous_jobs
        self.base_url = None
        self.output_folder = "outputs"
        self.job_title = None
        self.location = None
        self.job_id = None
        self.job_url = None
        self.data_dict = None
        self.secrets = None
        self.company_name = None
        self.cv_name_raw = None
        self.cover_letter_name_raw = None
        self.document = None
        self.template_path = None
        self.replacements = None
        self.cv_temp = None
        self.cv_path = None
        self.cv_name = None
        self.cover_letter_temp = None
        self.cover_letter_path = None
        self.cover_letter_name = None
        self.visited_jobs = None
        self.temporary_documents = []
        self.listings_array = []
        self.listings_iterator = None

    @staticmethod
    def return_clean_name(name):
        return name.lower().replace(" ", "_").replace("-", "_").replace("/", "_").replace("|", "_").replace("__", "_")
    
    def load_secrets(self):
        with open('secrets.json', 'r') as file:
            self.secrets = json.load(file)
    
    @staticmethod
    def wait_random():
        seconds = random.randint(1, 3)
        print(f"Waiting {seconds} seconds", end='\r')
        time.sleep(seconds)
    
    def log_in(self):

        email = self.secrets['seek_login']['email']
        password = self.secrets['seek_login']['password']
    
        text_box = WebDriverWait(self.driver, 10).until(
            EC.presence_of_element_located((By.ID, 'emailAddress'))
        )
        text_box.send_keys(email)
    
        text_box = WebDriverWait(self.driver, 5).until(
            EC.presence_of_element_located((By.ID, 'password'))
        )
        text_box.send_keys(password)
    
        login_button = WebDriverWait(self.driver, 5).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="signin_seekanz"]/div/div[4]/div/div[1]/button'))
        )
        self.wait_random()
        login_button.click()
        self.wait_random()

    def read_visited_jobs(self):
        with open('visited_jobs.json', 'r') as file:
            self.visited_jobs = json.load(file)
    
    def save_visited_job(self):
        self.visited_jobs[self.job_id] = self.data_dict
        with open('visited_jobs.json', 'w') as file:
            json.dump(self.visited_jobs, file, indent=4)

    def return_chat_gpt_response(self, query, job_description):
        client = OpenAI(api_key=self.secrets["chat_gpt_login"]["key"])
        message = f"{query}: {job_description}"
        chat_completion = client.chat.completions.create(
            messages=[{"role": "user", "content": message}], model="gpt-3.5-turbo",
        )
        return chat_completion.choices[0].message.content

    def set_document_name(self):
        self.document_name = self.return_clean_name(f"{self.document_type}_{self.name}_{self.job_title}.docx")
        
    def save_document(self):
        document_output_path = os.path.join(self.output_folder, self.document_name)
        self.document.save(document_output_path)

        document_temporary_path = os.path.join('temporary', self.document_name)
        shutil.copyfile(document_output_path, document_temporary_path)
        self.temporary_documents.append(document_temporary_path)

    def update_document(self):
        for key, value in self.replacements.items():
            for paragraph in self.document.paragraphs:
                if key in paragraph.text:
                    try:
                        runs = paragraph.runs
                        for i in range(len(runs)):
                            if key in runs[i].text:
                                text = runs[i].text.replace(key, value)
                                runs[i].text = text
                    except TypeError:
                        pass

                # backup
                if key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, value)

    def remove_temporary_files(self):
        print("Removing temporary files")
        for temporary_document in self.temporary_documents:
            try:
                os.remove(temporary_document)
            except:
                pass
        
    def click_quick_apply(self):
        quick_apply_button = WebDriverWait(self.driver, 2).until(
            EC.presence_of_element_located((By.LINK_TEXT, 'Quick apply'))
        )
        quick_apply_button.click()

    def load_pages(self, max_pages=1):
        self.base_url = r"https://www.seek.co.nz/api/chalice-search/v4/search?siteKey=NZ-Main&where=" + self.where + "&keywords=" + self.keywords
        for page_num in range(1, max_pages):
            self.wait_random()
            url = self.base_url + f"&page={page_num}"
            res = requests.get(url)
            print(res.status_code)
            self.listings_array.extend(res.json().get('data'))
        self.listings_iterator = iter(self.listings_array)

    def check_state(self):
        self.root.update()

        if self.skip_clicked:
            print(f"Skip clicked: {self.skip_clicked}")
            self.skip_clicked = False
            self.save_visited_job()
            self.remove_temporary_files()
            self.restart_thread()

        try:
            application_sent = WebDriverWait(self.driver, 1).until(EC.presence_of_element_located((By.ID, 'applicationSent')))
            self.messages_label.config(text=f"Application for {self.job_title} sent successfully.")
            self.remove_temporary_files()
            self.apply()
            self.data_dict["applied"] = True
        except:
            pass

        self.root.after(1000, self.check_state)

    def complete_application(self):
        if EC.presence_of_element_located((By.ID, 'password')):
            try:
                self.log_in()
            except:
                pass

        pyperclip.copy(self.temp_path_dict[self.document_type])
        self.messages_label.config(text="Please complete application.")


    def execute_chat_gpt_prompts(self):
        start_tag = '<CHAT_GPT>'
        end_tag = '</CHAT_GPT>'
        for paragraph in self.document.paragraphs:
            matches = re.findall(fr"{start_tag}(.*?){end_tag}", paragraph.text)
            if matches:
                for match in matches:
                    for tag in [start_tag, end_tag]:
                        match = match.replace(tag, "")

                    self.messages_label.config(text=f"Querying ChatGPT... ({match})")
                    query_response = self.return_chat_gpt_response(match, self.job_details)
                    placeholder_text = f"{start_tag}{match}{end_tag}"
                    paragraph.text = paragraph.text.replace(placeholder_text, query_response, 1)

    def apply(self):
        self.load_secrets()
        if self.driver is None:
            self.driver = webdriver.Chrome(service=ChromeService(), options=webdriver.ChromeOptions())

        try:
            listing = next(self.listings_iterator)
        except StopIteration:
            print("Completed all listings")
            self.driver.quit()
            return

        self.read_visited_jobs()
        self.temporary_documents = []
        self.company_name = listing.get("companyName")

        if self.company_name is None:
            self.company_name = listing.get("advertiser").get('description')

        self.job_title = listing.get("title")
        self.location = listing.get("location")
        self.job_id = listing.get('solMetadata').get('jobId')
        self.job_url = r"https://www.seek.co.nz/job/" + self.job_id

        self.data_dict = {
            "job_title": self.job_title,
            "company_name": self.company_name,
            "location": self.location,
            "url": self.job_url,
            "applied": False
        }

        if (self.job_id in self.visited_jobs.keys()) and (self.view_previous_jobs is True):
            if self.visited_jobs[self.job_id]['applied'] is True:
                print("Already applied")

        else:
            self.driver.get(self.job_url)
            self.messages_label.config(text=f"Got {self.job_url}")

            self.replacements = {
                "COMPANY_NAME": self.company_name,
                "JOB_TITLE": self.job_title,
                "LOCATION_NAME": self.location
            }

            self.job_details = self.driver.find_element(By.CSS_SELECTOR, 'div[data-automation="jobAdDetails"]').text


            for document_type in ['cv', 'cover_letter']:
                self.document_type = document_type
                self.template_path = self.temp_path_dict[document_type]
                self.document = Document(self.template_path)

                if document_type == 'cover_letter':
                    self.execute_chat_gpt_prompts()
                    self.update_document()

                self.set_document_name()
                self.save_document()

            try:
                self.click_quick_apply()

            except TimeoutException:
                self.messages_label.config(text="No quick apply")

            self.wait_random()
            self.complete_application()
            self.save_visited_job()

class Interface(Seeker):
    def __init__(self, root):
        self.root = root
        self.root.title("Seeker")
        self.driver = None
        self.temporary_documents = []
        self.temp_path_dict = {}

        self.main_frame = ttk.Frame(root)
        self.main_frame.pack(fill=tk.BOTH, expand=True, padx=40, pady=30)

        self.name_label = ttk.Label(self.main_frame, text="Your name:")
        self.name_label.grid(row=0, column=0, columnspan=2, pady=5, padx=5, sticky='E')
        self.name_entry = ttk.Entry(self.main_frame)
        self.name_entry.grid(row=0, column=2, columnspan=2, pady=5, padx=5, sticky='W')

        self.cv_template_label = ttk.Label(self.main_frame, text="CV Template:")
        self.cv_template_label.grid(row=2, column=0, columnspan=2, pady=5, padx=5, sticky='E')
        self.cv_template_button = ttk.Button(self.main_frame, text="Browse", command= lambda: self.browse_temp_path('cv'))
        self.cv_template_button.grid(row=2, column=2, columnspan=2, pady=5, padx=5, sticky='W')

        self.cover_letter_template_label = ttk.Label(self.main_frame, text="Cover Letter Template:")
        self.cover_letter_template_label.grid(row=3, column=0, columnspan=2, pady=5, padx=5, sticky='E')
        self.cover_letter_template_button = ttk.Button(self.main_frame, text="Browse", command= lambda: self.browse_temp_path('cover_letter'))
        self.cover_letter_template_button.grid(row=3, column=2, columnspan=2, pady=5, padx=5, sticky='W')

        self.where_label = ttk.Label(self.main_frame, text="Location:")
        self.where_label.grid(row=4, column=0, columnspan=2, pady=5, padx=5, sticky='E')
        self.where_entry = ttk.Entry(self.main_frame)
        self.where_entry.grid(row=4, column=2, columnspan=2, pady=5, padx=5, sticky='W')

        self.keywords_label = ttk.Label(self.main_frame, text="Keywords:")
        self.keywords_label.grid(row=5, column=0, columnspan=2, pady=5, padx=5, sticky='E')
        self.keywords_entry = ttk.Entry(self.main_frame)
        self.keywords_entry.grid(row=5, column=2, columnspan=2, pady=5, padx=5, sticky='W')

        self.view_previous_jobs_check = ttk.Checkbutton(self.main_frame, text='View previously-skipped jobs', variable=tk.StringVar(), onvalue=True, offvalue=False)
        self.view_previous_jobs_check.grid(row=6, column=0, columnspan=4, pady=5)
        self.view_previous_jobs_check.invoke()

        self.run_button = ttk.Button(self.main_frame, text="Run", command=self.run_action)
        self.run_button.grid(row=7, column=1, columnspan=2, pady=5)

        self.skip_button = ttk.Button(self.main_frame, text="Skip current listing", command=self.skip_action, state=DISABLED)
        self.skip_button.grid(row=8, column=1, columnspan=2, pady=5)

        self.quit_button = ttk.Button(self.main_frame, text="Quit", command=self.quit_action)
        self.quit_button.grid(row=9, column=1, columnspan=2, pady=5)

        self.messages_label = ttk.Label(self.main_frame, text="", wraplength=240, justify='center')
        self.messages_label.configure(anchor="center")
        self.messages_label.grid(row=10, column=0, columnspan=4, pady=5)

        self.temp_path_output = None
        self.thread = None
        self.skip_clicked = False

    def browse_temp_path(self, document_type):
        document_type_clean = document_type.replace("_", ' ')
        temp_path = filedialog.askopenfilename(title=f"Select temporary {document_type_clean} path")
        self.temp_path_dict[document_type] = temp_path
        self.messages_label.config(text=f"Set temporary {document_type_clean} path to {temp_path}")

    @staticmethod
    def return_clean_url_args(url_args):
        return url_args.replace(' ', '+')

    def start_thread(self):
        self.thread = threading.Thread(target=self.apply)
        self.thread.daemon = True
        self.thread.start()
        print("Thread started")

    def restart_thread(self):
        if self.thread is not None:
            self.thread.join()  # Wait for the thread to finish
            print("Thread stopped")
        self.start_thread()

    def skip_action(self):
        print(f"Skipping {self.job_title}")
        self.skip_clicked = True

    def quit_action(self):
        if self.driver is not None:
            self.driver.quit()
            self.thread.join()
        self.remove_temporary_files()
        sys.exit()

    @staticmethod
    def return_clean_text_input(text):
        return text.strip().replace(" ", "_")

    def run_action(self):
        self.check_state()
        self.run_button['state'] = 'disabled'
        self.name = self.return_clean_text_input(self.name_entry.get())
        self.temp_path = self.temp_path_output
        self.where = self.return_clean_url_args(self.where_entry.get())
        self.keywords = self.return_clean_url_args(self.keywords_entry.get())
        super().__init__(self.name, self.temp_path_dict, self.where, self.keywords, self.view_previous_jobs_check)

        self.load_pages(max_pages=2)
        self.skip_button['state'] = NORMAL
        self.messages_label.config(text="Running browser...")
        self.start_thread()


if __name__ == "__main__":
    root = tk.Tk()
    app = Interface(root)
    app.root.mainloop()

print("Complete")
